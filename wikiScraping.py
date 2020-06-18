#!/usr/bin/env python
# coding: utf-8

# In[76]:


from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException

import time
import json
import pandas as pd

import requests

import subprocess
import shutil

from docx import Document
import docx

import convertapi
convertapi.api_secret = 'your-api-secret'

#from docx2pdf import convert


# In[143]:


def getPage(keyword):
    # Path where you save the webdriver 
    executable_path = '/Users/charleschen/Downloads/geckodriver'

    # initiator the webdriver for Firefox browser
    driver = webdriver.Firefox(executable_path=executable_path)

    driver.implicitly_wait(10)

    # send a request
    #keywords = "DNA_microarray"
    driver.get("https://zh.wikipedia.org/wiki/%E9%A9%AC%E5%85%8B%E8%90%A8%E5%A7%86-%E5%90%89%E5%B0%94%E4%BC%AF%E7%89%B9%E6%B5%8B%E5%BA%8F")
    time.sleep(5)
    
    # Write Title
    title = driver.find_element_by_class_name("firstHeading").text
    mydoc.add_heading(title, 0)
    print(type(title))
    print(title)
    # get the whole page
    
    content = driver.find_elements_by_class_name("mw-content-ltr")[-1].text
#     print(content)
    mydoc.add_paragraph(content)
    
    
    docName = title + ".docx"
    pdfName = title + ".pdf"
    
    mydoc.save(docName)
    
    # save as pdf
    
#     file = open(docName, 'rb')
#     data = file.read()

#     input_filename = docName
#     output_filename = pdfName

#     p = subprocess.Popen(['unoconv', '--stdout', input_filename], stdout=subprocess.PIPE)
#     with open(output_filename, 'w') as output:
#         shutil.copyfileobj(p.stdout, output)
        
#     print("Finish Writing!")

    


# In[145]:


if __name__ == "__main__":
    mydoc = docx.Document()
    getPage("DNA_microarray")
    


# In[ ]:




